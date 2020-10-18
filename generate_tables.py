from docx.oxml import parse_xml
from docx import Document
from colors import *


def generate_tables(master, sort):

	document = Document()

	document.add_heading("Tableaux des stories")

	new_sort = {}

	for key, value in sort.items():

		new_sort[key] = {}

		table = document.add_table(0, len(value))
		cells = table.add_row().cells

		for i in range(len(value)):
			cells[0].merge(cells[i])
		cells[0].text = key
		cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))

		i = 1
		tmp = {}
		for c in value:
			t = f"{i} {c}"
			j = 1
			rows = []
			for u in value[c]:
				name = f"{i}.{j} {u['fields']['Name']}"
				row = u
				row['fields']["Name"] = name
				rows.append(row)
				j += 1
			tmp[t] = rows
			i += 1

		t = list(tmp)

		add = 1
		for kk in tmp:
			add = max(add, len(tmp[kk]))

		for u in range(add + 1):
			table.add_row()

		for u in range(len(t)):
			vv = tmp[t[u]]
			l = list(vv)

			table.cell(1, u).text = t[u]
			table.cell(1, u)._tc.get_or_add_tcPr().append(parse_xml(odd if u % 2 == 0 else even))

			for v in range(len(vv)):
				table.cell(v + 2, u).text = vv[v]["fields"]["Name"]
				table.cell(v + 2, u)._tc.get_or_add_tcPr().append(parse_xml(odd if u % 2 == 0 else even))

		document.add_paragraph()
		new_sort[key] = tmp

	document.add_page_break()
	master.append(document)

	return new_sort