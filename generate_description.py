from docx.oxml import parse_xml
from docx import Document
from colors import *

def generate_description(date, author):
	document = Document()

	document.add_heading("Description du document")

	table = document.add_table(0, 2)
	table.style = "TableGrid"

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Type de document\n").bold = True
	cells[1].text = "Project Log Document"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Date\n").bold = True
	cells[1].text = date
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(even))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Responsable du groupe\n").bold = True
	cells[1].text = "Hugo Frugier"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Auteur\n").bold = True
	cells[1].text = author
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(even))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Groupe\n").bold = True
	cells[1].text = "Hugo Frugier, Alexandre Tahery, Clément Chanal, Thomas Bouvier, Paul Gaston, Alexis Auriac, Alexandre Lefèvre"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Responsable de la relecture\n").bold = True
	cells[1].text = "Thomas Bouvier"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(even))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Mèl\n").bold = True
	cells[1].text = "mediwatch_2022@labeip.epitech.eu"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd))

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Sujet\n").bold = True
	cells[1].text = "Project Log Document du projet Mediwatch"
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(even))


	return document