from docx.oxml import parse_xml
from docx import Document
from colors import *

def generate_cards(title, content):

	document = Document()

	document.add_heading(title, 3)

	for subtitle in content:
		document.add_heading(subtitle, 4)

		for c in content[subtitle]:

			table = document.add_table(rows=0, cols=2)
			table.style = "TableGrid"

			try:
				cells = table.add_row().cells
				cells[0].merge(cells[1])
				cells[0]._tc.get_or_add_tcPr().append(parse_xml(odd))
				cells[0].text = c["fields"]["Name"].replace("\_", "_")
			except: pass

			try:
				cells = table.add_row().cells
				cells[0]._tc.get_or_add_tcPr().append(parse_xml(even))
				cells[1]._tc.get_or_add_tcPr().append(parse_xml(even))
				cells[0].text = "En tant que : {0}".format(c["fields"]["Wanter"])
				cells[1].text = "Je veux : {0}".format(c["fields"]["Wanted"])
			except: pass

			try:
				cells = table.add_row().cells
				cells[0].merge(cells[1])
				cells[0]._tc.get_or_add_tcPr().append(parse_xml(odd))
				cells[0].text = "Description :\n"
				cells[0].text += c["fields"]["Resumé"].strip().replace("\_", "_")
			except: pass

			try:
				cells = table.add_row().cells
				cells[0].merge(cells[1])
				cells[0]._tc.get_or_add_tcPr().append(parse_xml(even))
				cells[0].text = "Definition of Done :\n"
				cells[0].text += c["fields"]["Definition of done"].replace("\-", "-").replace("\_", "_")
			except: pass

			try:
				cells = table.add_row().cells
				cells[0]._tc.get_or_add_tcPr().append(parse_xml(odd))
				cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd))
				cells[0].text = "Charge estimée :"
				cells[1].text = f"{c['fields']['Duration']} j/homme"
			except: pass

			document.add_page_break()

	return document