from docx.oxml import parse_xml
from docx import Document
from colors import *
from strings import *

advancements = [
	{
		"date": START_DATE,
		"version": "0.1",
		"auteur": "PLD-BOT",
		"section": "Toutes",
		"commentaire": "Création du document"
	}
]

def generate_advancements(document):

	document.add_heading("Tableau des révisions")

	table = document.add_table(0, 5)
	table.style = "TableGrid"

	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("Date").bold = True
	cells[0]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[1].paragraphs[0].add_run("Version").bold = True
	cells[1]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[2].paragraphs[0].add_run("Auteur").bold = True
	cells[2]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[3].paragraphs[0].add_run("Section").bold = True
	cells[3]._tc.get_or_add_tcPr().append(parse_xml(title))
	cells[4].paragraphs[0].add_run("Commentaire").bold = True
	cells[4]._tc.get_or_add_tcPr().append(parse_xml(title))

	for i in range(len(advancements)):
		cells = table.add_row().cells

		cells[0].paragraphs[0].text = advancements[i]["date"]
		cells[0]._tc.get_or_add_tcPr().append(parse_xml(odd if i % 2 == 0 else even))
		cells[1].paragraphs[0].text = advancements[i]["version"]
		cells[1]._tc.get_or_add_tcPr().append(parse_xml(odd if i % 2 == 0 else even))
		cells[2].paragraphs[0].text = advancements[i]["auteur"]
		cells[2]._tc.get_or_add_tcPr().append(parse_xml(odd if i % 2 == 0 else even))
		cells[3].paragraphs[0].text = advancements[i]["section"]
		cells[3]._tc.get_or_add_tcPr().append(parse_xml(odd if i % 2 == 0 else even))
		cells[4].paragraphs[0].text = advancements[i]["commentaire"]
		cells[4]._tc.get_or_add_tcPr().append(parse_xml(odd if i % 2 == 0 else even))

	document.add_page_break()

	return document