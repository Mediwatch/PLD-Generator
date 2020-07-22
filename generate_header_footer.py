from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def __create_element(name):
    return OxmlElement(name)


def __create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def __add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = __create_element('w:t')
    __create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = __create_element('w:fldChar')
    __create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = __create_element('w:instrText')
    __create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "Page "

    fldChar2 = __create_element('w:fldChar')
    __create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = __create_element('w:t')
    __create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' sur '
    of_run._r.append(t2)

    fldChar3 = __create_element('w:fldChar')
    __create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = __create_element('w:instrText')
    __create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = __create_element('w:fldChar')
    __create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


def generate_header_footer(document):
	document.sections[0].different_first_page_header_footer = True

	table = document.sections[0].header.add_table(1, 2, Inches(6))
	table.cell(0, 0).add_paragraph().add_run().add_picture("MediwatchLogo.png", Inches(1))
	table.cell(0, 1).add_paragraph().add_run().add_picture("EIPLogo.png", Inches(2))

	__add_page_number(document.sections[0].footer.paragraphs[0])