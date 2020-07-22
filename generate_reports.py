from docx import Document
from datetime import datetime
from operator import itemgetter
import re

from pprint import pp

def generate_reports(reports):

	document = Document()

	r = []

	for report in reports:
		try:
			date = datetime.strptime(report['fields']['Date'], "%Y-%m-%d").strftime("%d/%m/%Y")
			r.append({
				"content": report['fields']['Rapport'],
				"title": f"{date} - Rapport d'avancement pour le {report['fields']['Type']} du {report['fields']['Sprint']} sprint"
			})
		except: pass

	def sort(l):
		new_list = []

		for i in l:
			if len(new_list) == 0:
				new_list.append(i)
				continue

			curr_date = datetime.strptime(i["title"][:i["title"].index(" - ")], "%d/%m/%Y")
			for idx in range(len(new_list)):
				new_r = new_list[idx]
				new_date = datetime.strptime(new_r["title"][:i["title"].index(" - ")], "%d/%m/%Y")
				if curr_date > new_date:
					new_list.insert(idx, i)
					break

		return new_list

	r = sort(r)

	document.add_heading("Rapports d'avancement")

	for report in r:
	 	document.add_heading(report["title"], 3)
	 	document.add_paragraph(report["content"])

	document.add_page_break()

	return document