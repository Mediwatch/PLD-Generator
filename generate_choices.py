from docx.oxml import parse_xml
from docx import Document
from colors import *

from pprint import pp

def generate_choices(sort, individuals):

	document = Document()

	document.add_heading("Répartition")

	group = []

	total = 0
	for s in sort.values():
		for ss in s.values():
			for sss in ss:
				total += int(sss["fields"]["Duration"])
				group.append(sss["fields"]["Associate to"]["name"])
	group = list(set(group))

	document.add_paragraph(f"Au total, le sprint totalise {total} j/homme.")
	document.add_paragraph(f"Le sprint est réparti sur {len(sort)} grande{'' if len(sort) == 1 else 's'} partie{'' if len(sort) == 1 else 's'} :")

	s = ""
	for i in sort.items():
		total = 0
		assigned = []
		for ii in i[1].values():
			for iii in ii:
				total += int(iii["fields"]["Duration"])
				assigned.append(iii["fields"]["Associate to"]["name"])

		assigned = list(set(assigned))
		s += f"\t- {i[0]} : {total} j/homme "
		s += "(" + ", ".join(assigned) + ")"
		s += "\n"

	document.add_paragraph(s)

	# rapport personnels

	for report in individuals:
		document.add_heading(f"Rapport de {report['fields']['Nom']['name']}", 4)
		try:
			document.add_paragraph(report["fields"]["Résumé"])
		except: pass

	document.add_page_break()

	return document