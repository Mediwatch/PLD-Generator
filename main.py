from docx import Document
from docxcompose.composer import Composer

from win32com import client
import inspect, os

from strings import *

from get_airtable import get_airtable
from generate_header_footer import generate_header_footer
from generate_chart import generate_chart
from generate_description import generate_description
from generate_advancements import generate_advancements
from generate_tables import generate_tables
from generate_choices import generate_choices
from generate_cards import generate_cards
from generate_reports import generate_reports
from generate_archives import generate_archives

def create_doc_from_file(file, line_break=True):
	d = Document(file)
	if line_break:
		d.add_page_break()
	return d

def separate_cards(content, archives, cards, sort):
	for c in content:
		try:
			archives.append(c) if c["fields"]["Status"].startswith("Archived") else cards.append(c)
		except:
			pass
	for c in cards:
		key, value = c["fields"]["Task Type"], c
		if key not in sort.keys():
			sort[key] = {value["fields"]["TypeTitle"]: [c]}
		else:
			if value["fields"]["TypeTitle"] not in sort[key]:
				sort[key][value["fields"]["TypeTitle"]] = [c]
			else:
				sort[key][value["fields"]["TypeTitle"]].append(c)


def create_stories(master, tmp):
	tmp = Document()
	tmp.add_heading("Stories des livrables")
	master.append(tmp)
	for key in sort:
		master.append(generate_cards(key, sort[key]))

airtable = get_airtable()
content = airtable[0]
reports = airtable[1]
individuals = airtable[2]
sort = {}
archives = []
cards = []

document = Document()
master = Composer(document)
cover_abstract = create_doc_from_file("template_cover_abstract.docx")
summary = create_doc_from_file("template_summary.docx")
chart = create_doc_from_file("template_chart.docx", False)
generate_header_footer(document)
separate_cards(content, archives, cards, sort)
generate_chart(chart, sort)

master.append(cover_abstract)
master.append(generate_advancements(generate_description(START_DATE, "PLD-Bot")))
master.append(summary)
master.append(chart)
master.append(generate_choices(sort, individuals))
sort = generate_tables(master, sort)
create_stories(master, sort)
master.append(generate_reports(reports))
master.append(generate_archives(archives))

master.save(PLD_NAME)

# Update ToC https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux
# Windows only
script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
word = client.DispatchEx("Word.Application")
doc = word.Documents.Open(os.path.join(script_dir, PLD_NAME))
doc.TablesOfContents(1).Update()
doc.Close(SaveChanges=True)
word.Quit()

os.startfile(PLD_NAME)